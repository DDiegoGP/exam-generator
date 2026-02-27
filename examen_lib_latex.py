import pandas as pd
import random, re, os
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl

try:
    import gspread
    from google.auth import default
    _HAS_GSPREAD = True
except ImportError:
    _HAS_GSPREAD = False

# --- DETECCIÓN DE ENTORNO ---
def detectar_entorno():
    """Retorna 'colab' si se ejecuta en Google Colab, 'local' en caso contrario."""
    try:
        import google.colab
        return 'colab'
    except ImportError:
        return 'local'

# --- CONEXIÓN GOOGLE SHEETS ---
def conectar_sheet_seguro(url):
    if not _HAS_GSPREAD:
        raise ImportError("gspread no disponible. Instale gspread y google-auth para usar Google Sheets.")
    creds, _ = default()
    gc = gspread.authorize(creds)
    return gc.open_by_url(url)

# --- CONEXIÓN EXCEL LOCAL ---
def cargar_excel_local(filepath):
    """Lee todas las hojas de un .xlsx y devuelve dict {nombre_hoja: DataFrame}."""
    xls = pd.ExcelFile(filepath, engine='openpyxl')
    return {name: pd.read_excel(xls, sheet_name=name) for name in xls.sheet_names}

def backup_excel(filepath, max_backups=10):
    """Crea backup del Excel antes de guardar. Mantiene ultimas max_backups copias."""
    import shutil, datetime as _dt
    if not os.path.exists(filepath): return
    backup_dir = os.path.join(os.path.dirname(filepath), '_backups')
    os.makedirs(backup_dir, exist_ok=True)
    base = os.path.splitext(os.path.basename(filepath))[0]
    ts = _dt.datetime.now().strftime('%Y%m%d_%H%M%S')
    shutil.copy2(filepath, os.path.join(backup_dir, f"{base}_{ts}.xlsx"))
    # Limpiar backups antiguos
    bks = sorted([f for f in os.listdir(backup_dir) if f.startswith(base) and f.endswith('.xlsx')])
    for old in bks[:-max_backups]:
        os.remove(os.path.join(backup_dir, old))

def guardar_excel_local(filepath, dict_of_dfs):
    """Escribe a disco. Si filepath está vacío (modo cloud/upload), no hace nada."""
    if not filepath:
        return
    backup_excel(filepath)
    with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
        for sheet_name, df in dict_of_dfs.items():
            df.to_excel(writer, sheet_name=sheet_name, index=False)

def generar_excel_bytes(dict_of_dfs) -> bytes:
    """Genera el Excel en memoria (para descarga en cloud mode o backup)."""
    import io as _io
    buf = _io.BytesIO()
    with pd.ExcelWriter(buf, engine='openpyxl') as writer:
        for sheet_name, df in dict_of_dfs.items():
            df.to_excel(writer, sheet_name=sheet_name, index=False)
    return buf.getvalue()

def actualizar_pregunta_excel_local(filepath, dict_of_dfs, pid, datos):
    """Equivalente local de actualizar_pregunta_db para archivos Excel."""
    try:
        bloque = datos['bloque']
        if bloque not in dict_of_dfs:
            return False, f"Bloque '{bloque}' no encontrado"
        df = dict_of_dfs[bloque]
        # Asegurar columna notas existe
        if not any('nota' in str(c).lower() for c in df.columns):
            df['notas'] = ''; dict_of_dfs[bloque] = df
        mask = df['ID_Pregunta'].astype(str) == str(pid)
        if not mask.any():
            return False, "ID no encontrado"
        idx = df[mask].index[0]

        for col in df.columns:
            cl = str(col).lower().strip()
            if 'enunciado' in cl:
                df.at[idx, col] = datos['enunciado']
            elif 'tema' in cl and 'id' not in cl:
                df.at[idx, col] = datos['tema']
            elif 'correcta' in cl or 'resp' in cl:
                df.at[idx, col] = datos['correcta']
            elif 'dificultad' in cl:
                df.at[idx, col] = datos['dificultad']
            elif 'usada' in cl or 'used' in cl:
                if datos.get('usada'):
                    df.at[idx, col] = datos['usada']
            elif 'nota' in cl:
                df.at[idx, col] = datos.get('notas', '')

        # Actualizar opciones (columnas después de Enunciado)
        enun_col_idx = None
        for i, col in enumerate(df.columns):
            if 'enunciado' in str(col).lower():
                enun_col_idx = i
                break
        if enun_col_idx is not None:
            for j, op in enumerate(datos['opciones']):
                op_col_idx = enun_col_idx + 1 + j
                if op_col_idx < len(df.columns):
                    df.iat[idx, op_col_idx] = op

        guardar_excel_local(filepath, dict_of_dfs)
        return True, "Actualizado correctamente"
    except Exception as e:
        return False, str(e)

def eliminar_preguntas_excel_local(filepath, dict_of_dfs, ids):
    """Elimina varias preguntas por ID de todas las hojas del Excel."""
    try:
        backup_excel(filepath)
        ids_set = set(str(i) for i in ids)
        count = 0
        for sheet in list(dict_of_dfs.keys()):
            df = dict_of_dfs[sheet]
            if 'ID_Pregunta' not in df.columns:
                continue
            mask = df['ID_Pregunta'].astype(str).isin(ids_set)
            count += int(mask.sum())
            dict_of_dfs[sheet] = df[~mask].reset_index(drop=True)
        guardar_excel_local(filepath, dict_of_dfs)
        return True, f"{count} pregunta(s) eliminada(s)"
    except Exception as e:
        return False, str(e)

def actualizar_campo_masivo(filepath, dict_of_dfs, ids, campo, valor):
    """Actualiza tema o dificultad para múltiples preguntas a la vez."""
    try:
        backup_excel(filepath)
        ids_set = set(str(i) for i in ids)
        count = 0
        for sheet, df in dict_of_dfs.items():
            if 'ID_Pregunta' not in df.columns:
                continue
            mask = df['ID_Pregunta'].astype(str).isin(ids_set)
            if not mask.any():
                continue
            for col in df.columns:
                cl = str(col).lower().strip()
                if campo == 'tema' and 'tema' in cl and 'id' not in cl:
                    dict_of_dfs[sheet].loc[mask, col] = valor
                    count += int(mask.sum())
                elif campo == 'dificultad' and 'dificultad' in cl:
                    dict_of_dfs[sheet].loc[mask, col] = valor
                    count += int(mask.sum())
        guardar_excel_local(filepath, dict_of_dfs)
        return True, f"{count} pregunta(s) actualizadas → {campo}='{valor}'"
    except Exception as e:
        return False, str(e)

def reemplazar_texto_masivo(filepath, dict_of_dfs, ids, buscar, reemplazar_con):
    """Find & replace en el enunciado de las preguntas indicadas."""
    try:
        backup_excel(filepath)
        ids_set = set(str(i) for i in ids)
        count = 0
        for sheet, df in dict_of_dfs.items():
            if 'ID_Pregunta' not in df.columns:
                continue
            mask = df['ID_Pregunta'].astype(str).isin(ids_set)
            if not mask.any():
                continue
            for col in df.columns:
                if 'enunciado' in str(col).lower():
                    dict_of_dfs[sheet].loc[mask, col] = (
                        df.loc[mask, col].astype(str)
                        .str.replace(buscar, reemplazar_con, regex=False)
                    )
                    count += int(mask.sum())
                    break
        guardar_excel_local(filepath, dict_of_dfs)
        return True, f"Texto reemplazado en {count} pregunta(s)"
    except Exception as e:
        return False, str(e)

def validar_pregunta(pregunta):
    """Valida una pregunta dict. Retorna (ok, warnings_list)."""
    warnings = []
    enun = pregunta.get('enunciado', '').strip()
    if not enun: warnings.append("Enunciado vacio")
    ops = pregunta.get('opciones_list', [])
    if len(ops) < 4: warnings.append(f"Solo {len(ops)} opciones (se necesitan 4)")
    for i, op in enumerate(ops):
        if not str(op).strip(): warnings.append(f"Opcion {chr(65+i)} vacia")
    correcta = pregunta.get('letra_correcta', '').upper()
    if correcta not in ('A','B','C','D'): warnings.append(f"Respuesta correcta invalida: '{correcta}'")
    ops_clean = [str(o).strip().lower() for o in ops if str(o).strip()]
    if len(ops_clean) != len(set(ops_clean)): warnings.append("Opciones duplicadas detectadas")
    return len(warnings) == 0, warnings

# --- UTILIDADES ---
def check_for_similar_enunciado(text, df):
    if df.empty: return False, 0.0
    from difflib import SequenceMatcher
    text = str(text).lower().strip()
    if 'enunciado' not in df.columns: return False, 0.0
    sims = df['enunciado'].astype(str).apply(lambda x: SequenceMatcher(None, text, x.lower().strip()).ratio())
    max_sim = sims.max()
    return (max_sim > 0.9, max_sim)

def generar_siguiente_id(df, bloque, tema):
    prefix = f"FM_{str(bloque).zfill(2)}_{str(tema).zfill(2)}"
    if df is None or df.empty or 'ID_Pregunta' not in df.columns: return f"{prefix}_01", 1
    existing = df[df['ID_Pregunta'].astype(str).str.startswith(prefix)]
    if existing.empty: return f"{prefix}_01", 1
    nums = []
    for x in existing['ID_Pregunta']:
        try: nums.append(int(x.split('_')[-1]))
        except: pass
    next_n = max(nums) + 1 if nums else 1
    return f"{prefix}_{str(next_n).zfill(2)}", next_n

def get_first_empty_row(worksheet, col_check=1):
    cols = worksheet.col_values(col_check) 
    return len(cols) + 1

# --- PARSERS DE IMPORTACIÓN (formato normalizado) ---
# Cada parser devuelve lista de dicts:
# {'enunciado': str, 'opciones_list': [A,B,C,D], 'letra_correcta': str,
#  'bloque': str, 'tema': str, 'dificultad': str, '_warnings': [str]}

def parse_aiken(text, bloque_destino='', tema_destino='1', dificultad_destino='Media'):
    """Parsea texto en formato Aiken. Devuelve lista normalizada de preguntas."""
    preguntas = []
    lines = [l.rstrip() for l in text.splitlines()]
    current = None
    re_opt = re.compile(r'^([A-D])[.)]\s+(.+)', re.IGNORECASE)
    re_ans = re.compile(r'^ANSWER[:\s]+([A-D])', re.IGNORECASE)

    for line in lines:
        if not line.strip():
            continue
        m_ans = re_ans.match(line)
        m_opt = re_opt.match(line)
        if m_ans:
            if current:
                current['letra_correcta'] = m_ans.group(1).upper()
                while len(current['opciones_list']) < 4:
                    current['opciones_list'].append("")
                current['opciones_list'] = current['opciones_list'][:4]
                preguntas.append(current)
                current = None
        elif m_opt:
            if current:
                current['opciones_list'].append(m_opt.group(2).strip())
        else:
            # Nueva pregunta: línea que no es opción ni ANSWER
            if current:  # pregunta sin ANSWER: guardar incompleta
                current['_warnings'].append("Sin respuesta ANSWER")
                while len(current['opciones_list']) < 4: current['opciones_list'].append("")
                current['opciones_list'] = current['opciones_list'][:4]
                preguntas.append(current)
            current = {
                'enunciado': line.strip(), 'opciones_list': [],
                'letra_correcta': 'A', 'bloque': bloque_destino,
                'tema': tema_destino, 'dificultad': dificultad_destino,
                '_warnings': []
            }
    if current:
        current['_warnings'].append("Sin respuesta ANSWER al final")
        while len(current['opciones_list']) < 4: current['opciones_list'].append("")
        current['opciones_list'] = current['opciones_list'][:4]
        preguntas.append(current)

    # Validación básica
    for p in preguntas:
        if len([o for o in p['opciones_list'] if o.strip()]) < 4:
            p['_warnings'].append("Menos de 4 opciones detectadas")
    return preguntas

# --- IMPORTACIÓN WORD / PDF ---
MARCAS_CORRECTA_WORD = [
    "Negrita", "Resaltado (highlight)", "Color (cualquiera)",
    "Subrayado", "Asterisco (*)", "MAYÚSCULAS", "Siempre la primera"
]

def procesar_archivo_docx(filepath, bloque_destino, tema_destino='1', dificultad_destino='Media', marca_correcta='Negrita'):
    """Parsea un .docx con preguntas numeradas. Devuelve lista normalizada."""
    from docx.enum.text import WD_COLOR_INDEX
    doc = Document(filepath)
    preguntas = []
    current_preg = None
    re_preg = re.compile(r'^\s*(\d+)[\.\-\)]+\s*(.+)', re.DOTALL)
    re_opt_explicit = re.compile(r'^\s*([a-dA-D])[\.\-\)]+\s*(.+)', re.IGNORECASE)

    def _es_correcta(para, texto):
        if marca_correcta == 'Negrita':
            return any(r.bold for r in para.runs if r.text.strip())
        if marca_correcta == 'Resaltado (highlight)':
            for r in para.runs:
                try:
                    hc = r.font.highlight_color
                    if hc is not None and hc not in (WD_COLOR_INDEX.AUTO, WD_COLOR_INDEX.WHITE, None):
                        return True
                except Exception:
                    pass
            return False
        if marca_correcta == 'Color (cualquiera)':
            for r in para.runs:
                try:
                    if r.font.color.type is not None and r.font.color.rgb != RGBColor(0, 0, 0):
                        return True
                except Exception:
                    pass
            return False
        if marca_correcta == 'Subrayado':
            return any(r.underline for r in para.runs if r.text.strip())
        if marca_correcta == 'Asterisco (*)':
            t = texto.strip()
            return t.startswith('*') or t.endswith('*')
        if marca_correcta == 'MAYÚSCULAS':
            t = texto.strip()
            return t == t.upper() and any(c.isalpha() for c in t)
        return False  # 'Siempre la primera' → default A

    # Estilos de sección/título que hay que ignorar
    _HEADING_STYLES = {
        'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5',
        'Title', 'Subtitle', 'TOC Heading',
    }

    for para in doc.paragraphs:
        txt = para.text.strip()
        if not txt:
            continue
        # Saltar encabezados con estilo Word de título/sección
        style_name = (getattr(para.style, 'name', '') or '').strip()
        if style_name in _HEADING_STYLES:
            continue
        # Saltar líneas sospechosas de ser encabezados/pies de página:
        # muy cortas (y que no sean una opción a/b/c/d), o solo números
        if len(txt) < 5 and not re_opt_explicit.match(txt):
            continue

        m_p = re_preg.match(txt)
        if m_p:
            # Enunciado tiene que tener algo de sustancia
            enun_candidato = m_p.group(2).strip()
            if len(enun_candidato) < 8:
                continue
            if current_preg:
                while len(current_preg['opciones_list']) < 4: current_preg['opciones_list'].append("")
                current_preg['opciones_list'] = current_preg['opciones_list'][:4]
                preguntas.append(current_preg)
            current_preg = {
                'enunciado': enun_candidato, 'opciones_list': [],
                'letra_correcta': 'A', 'bloque': bloque_destino,
                'tema': tema_destino, 'dificultad': dificultad_destino,
                '_warnings': []
            }
            continue
        if current_preg:
            m_o = re_opt_explicit.match(txt)
            texto_op = m_o.group(2).strip() if m_o else (txt if len(current_preg['opciones_list']) < 4 else "")
            # Limpiar marcador asterisco del texto si aplica
            clean_op = texto_op.strip().strip('*').strip() if marca_correcta == 'Asterisco (*)' else texto_op
            if clean_op:
                current_preg['opciones_list'].append(clean_op)
                if _es_correcta(para, texto_op):
                    idx = len(current_preg['opciones_list']) - 1
                    current_preg['letra_correcta'] = ['A', 'B', 'C', 'D'][idx] if idx < 4 else 'A'

    if current_preg:
        while len(current_preg['opciones_list']) < 4: current_preg['opciones_list'].append("")
        current_preg['opciones_list'] = current_preg['opciones_list'][:4]
        preguntas.append(current_preg)

    for p in preguntas:
        if len([o for o in p['opciones_list'] if o.strip()]) < 4:
            p['_warnings'].append("Menos de 4 opciones detectadas")
    return preguntas

def parse_pdf_bytes(pdf_bytes, bloque_destino='', tema_destino='1', dificultad_destino='Media', marca_correcta='Negrita'):
    """Parsea un PDF digital (bytes) con preguntas numeradas. Usa pdfplumber.
    Detecta: preguntas numeradas (1. / 1) / 1-) y opciones (a) / A. / A-)
    Para 'Negrita': intenta detectar fontname con 'Bold'.
    Para 'Asterisco (*)' y 'MAYÚSCULAS': igual que Word.
    Para el resto de marcas o si no se detecta: la correcta queda como A (editar manualmente).
    """
    import io as _io
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("Instala pdfplumber: pip install pdfplumber")

    re_preg    = re.compile(r'^\s*(\d+)[\.\-\)]+\s*(.+)')
    re_opt     = re.compile(r'^\s*([a-dA-D])[\.\-\)]+\s*(.+)', re.IGNORECASE)
    re_ans     = re.compile(r'^ANSWER[:\s]+([A-D])', re.IGNORECASE)
    re_pagnum  = re.compile(r'^(pág(ina)?\.?|page?\.?|p\.)\s*\d+', re.IGNORECASE)

    def _extract_page_lines(page):
        """Devuelve lista de (texto, is_bold) recortando el área de encabezado y pie."""
        # Recortar ~7% superior e inferior para saltar encabezados/pies de página
        margin_v = page.height * 0.07
        try:
            body = page.within_bbox((0, margin_v, page.width, page.height - margin_v))
        except Exception:
            body = page
        chars = body.chars or []
        lines = []
        if not chars:
            txt = body.extract_text() or ""
            return [(l.rstrip(), False) for l in txt.splitlines()]
        chars_sorted = sorted(chars, key=lambda c: (round(c['top'], 1), c['x0']))
        cur_y, cur_text, cur_bold = None, "", False
        for ch in chars_sorted:
            y = round(ch['top'], 1)
            if cur_y is None:
                cur_y = y
            if abs(y - cur_y) > 3:
                lines.append((cur_text.rstrip(), cur_bold))
                cur_text, cur_bold, cur_y = "", False, y
            cur_text += ch.get('text', '')
            if 'Bold' in (ch.get('fontname') or ''):
                cur_bold = True
        if cur_text.strip():
            lines.append((cur_text.rstrip(), cur_bold))
        return lines

    raw_lines = []
    with pdfplumber.open(_io.BytesIO(pdf_bytes)) as pdf:
        for page in pdf.pages:
            raw_lines.extend(_extract_page_lines(page))

    preguntas = []
    current   = None
    for line_txt, line_bold in raw_lines:
        txt = line_txt.strip()
        if not txt:
            continue
        # Filtrar líneas que probablemente son encabezados o pies de página
        if len(txt) < 5 and not re_opt.match(txt):   # muy corta y no es opción
            continue
        if re_pagnum.match(txt):                       # patrón "Página X / Page X"
            continue
        if txt.isdigit():                              # número solo (nº de página)
            continue
        m_ans  = re_ans.match(txt)
        m_preg = re_preg.match(txt)
        m_opt  = re_opt.match(txt)

        if m_ans:
            if current:
                current['letra_correcta'] = m_ans.group(1).upper()
                while len(current['opciones_list']) < 4: current['opciones_list'].append("")
                current['opciones_list'] = current['opciones_list'][:4]
                preguntas.append(current); current = None
        elif m_preg and not m_opt:
            enun_cand = m_preg.group(2).strip()
            if len(enun_cand) < 8:   # demasiado corto para ser una pregunta real
                continue
            if current:
                while len(current['opciones_list']) < 4: current['opciones_list'].append("")
                current['opciones_list'] = current['opciones_list'][:4]
                preguntas.append(current)
            current = {
                'enunciado': enun_cand, 'opciones_list': [],
                'letra_correcta': 'A', 'bloque': bloque_destino,
                'tema': tema_destino, 'dificultad': dificultad_destino, '_warnings': []
            }
        elif m_opt and current:
            texto_op = m_opt.group(2).strip()
            clean_op = texto_op.strip().strip('*').strip() if marca_correcta == 'Asterisco (*)' else texto_op
            if clean_op:
                current['opciones_list'].append(clean_op)
                idx = len(current['opciones_list']) - 1
                es_c = False
                if   marca_correcta == 'Negrita'       and line_bold:           es_c = True
                elif marca_correcta == 'Asterisco (*)'  and (texto_op.startswith('*') or texto_op.endswith('*')): es_c = True
                elif marca_correcta == 'MAYÚSCULAS'    and texto_op == texto_op.upper() and any(c.isalpha() for c in texto_op): es_c = True
                if es_c:
                    current['letra_correcta'] = ['A','B','C','D'][idx] if idx < 4 else 'A'
        elif current and not m_preg:
            # Línea de continuación del enunciado
            if not current['opciones_list']:
                current['enunciado'] += ' ' + txt

    if current:
        while len(current['opciones_list']) < 4: current['opciones_list'].append("")
        current['opciones_list'] = current['opciones_list'][:4]
        preguntas.append(current)

    for p in preguntas:
        if len([o for o in p['opciones_list'] if o.strip()]) < 4:
            p['_warnings'].append("Menos de 4 opciones detectadas")
        if p['letra_correcta'] == 'A' and marca_correcta not in ('Siempre la primera', 'Asterisco (*)', 'MAYÚSCULAS'):
            p['_warnings'].append("Respuesta A por defecto — verifica")
    return preguntas


# --- DB UPDATE ---
def actualizar_pregunta_db(sheet, pid, datos, idx_usada_ignorado):
    try:
        ws = sheet.worksheet(datos['bloque'])
        cell = ws.find(pid)
        if not cell: return False, "ID no encontrado"
        r = cell.row
        headers = [str(h).lower().strip() for h in ws.row_values(1)]
        col_enun = -1; col_tem = -1; col_corr = -1; col_dif = -1; col_usada = -1; col_nota = -1
        for i, h in enumerate(headers):
            idx = i + 1
            if 'enunciado' in h: col_enun = idx
            elif 'tema' in h: col_tem = idx
            elif 'correcta' in h or 'resp' in h: col_corr = idx
            elif 'dificultad' in h: col_dif = idx
            elif 'usada' in h or 'used' in h: col_usada = idx
            elif 'nota' in h: col_nota = idx

        if col_enun != -1: ws.update_cell(r, col_enun, datos['enunciado'])
        if col_tem != -1: ws.update_cell(r, col_tem, datos['tema'])
        if col_corr != -1: ws.update_cell(r, col_corr, datos['correcta'])
        if col_dif != -1: ws.update_cell(r, col_dif, datos['dificultad'])
        if col_nota != -1: ws.update_cell(r, col_nota, datos.get('notas', ''))
        if col_enun != -1:
            cells = [gspread.Cell(r, col_enun+1+i, op) for i, op in enumerate(datos['opciones'])]
            ws.update_cells(cells)
        if datos['usada'] and col_usada != -1:
            ws.update_cell(r, col_usada, datos['usada'])
        return True, "Actualizado correctamente"
    except Exception as e: return False, str(e)

# ── Esquemas de color para Word (R,G,B) ────────────────────────────────────
_WORD_COLOR_SCHEMES = {
    'azul': {'primary': (41, 128, 185), 'secondary': (52, 73, 94),  'bg': (236, 240, 241)},
    'ucm':  {'primary': (165, 28, 48),  'secondary': (120, 20, 35), 'bg': (245, 240, 235)},
    'byn':  {'primary': (60, 60, 60),   'secondary': (40, 40, 40),  'bg': (245, 245, 245)},
}

# ── Paletas de color para el .sty (LaTeX) ──────────────────────────────────
_STY_COLORS = {
    'azul': (
        r'\definecolor{primario}{RGB}{41, 128, 185}'  '\n'
        r'\definecolor{secundario}{RGB}{52, 73, 94}'  '\n'
        r'\definecolor{acento}{RGB}{231, 76, 60}'     '\n'
        r'\definecolor{fondo}{RGB}{236, 240, 241}'    '\n'
        r'\definecolor{texto}{RGB}{44, 62, 80}'
    ),
    'ucm': (
        r'\definecolor{primario}{RGB}{165, 28, 48}'   '\n'
        r'\definecolor{secundario}{RGB}{120, 20, 35}' '\n'
        r'\definecolor{acento}{RGB}{184, 134, 11}'    '\n'
        r'\definecolor{fondo}{RGB}{245, 240, 235}'    '\n'
        r'\definecolor{texto}{RGB}{40, 40, 40}'
    ),
    'byn': (
        r'\definecolor{primario}{RGB}{60, 60, 60}'    '\n'
        r'\definecolor{secundario}{RGB}{40, 40, 40}'  '\n'
        r'\definecolor{acento}{RGB}{100, 100, 100}'   '\n'
        r'\definecolor{fondo}{RGB}{245, 245, 245}'    '\n'
        r'\definecolor{texto}{RGB}{0, 0, 0}'
    ),
}

# ── Fuentes LaTeX (comando usepackage a inyectar) ───────────────────────────
_STY_FONTS = {
    'cm':        '',
    'palatino':  r'\usepackage{mathpazo}',
    'times':     r'\usepackage{mathptmx}',
    'libertine': r'\usepackage{libertine}' '\n' r'\usepackage[libertine]{newtxmath}',
    'helvet':    r'\usepackage{helvet}' '\n' r'\renewcommand{\familydefault}{\sfdefault}',
    'garamond':  r'\usepackage{garamondx}',
}

# ── Fuentes Word (nombre de fuente) ────────────────────────────────────────
_WORD_FONTS = {
    'cm':        'Cambria',
    'palatino':  'Palatino Linotype',
    'times':     'Times New Roman',
    'libertine': 'Cambria',
    'helvet':    'Arial',
    'garamond':  'Garamond',
}


def _generar_sty(cfg) -> bytes:
    """Carga el .sty y aplica overrides de color/fuente/modo según cfg."""
    import os as _os
    sty_path = _os.path.join(_os.path.dirname(_os.path.abspath(__file__)),
                              'estilo_examen_moderno_v2.sty')
    try:
        with open(sty_path, 'r', encoding='utf-8') as _f:
            sty = _f.read()
    except FileNotFoundError:
        return b""

    overrides = []
    scheme = cfg.get('color_scheme', 'azul')
    if scheme in _STY_COLORS and scheme != 'azul':
        overrides.append('% === Esquema de color seleccionado ===')
        overrides.append(_STY_COLORS[scheme])
    font = cfg.get('tipografia', 'cm')
    if font in _STY_FONTS and _STY_FONTS[font]:
        overrides.append('% === Tipografía seleccionada ===')
        overrides.append(_STY_FONTS[font])
    if cfg.get('modo_compacto', False):
        overrides.append(r'\def\modocompacto{1}')

    if overrides:
        # Inyectar DESPUÉS de \logo{} (los \ifdefined\modocompacto vienen después)
        inject_block = r'\logo{}' + '\n' + '\n'.join(overrides) + '\n'
        sty = sty.replace(r'\logo{}', inject_block, 1)

    return sty.encode('utf-8')


def _gen_caja_alumno_latex(campos: list) -> str:
    """Genera tcolorbox con campos del alumno según selección."""
    if not campos:
        return ''
    lines = []
    if 'nombre' in campos:
        lines.append(r'\textbf{Nombre y Apellidos:} \hrulefill')
    row2 = []
    if 'dni'   in campos:
        row2.append(r'\begin{minipage}[t]{0.45\textwidth}\textbf{DNI/NIU:} \hrulefill\end{minipage}')
    if 'grupo' in campos:
        row2.append(r'\begin{minipage}[t]{0.45\textwidth}\textbf{Grupo:} \hrulefill\end{minipage}')
    if row2:
        lines.append(r'\\[8mm]' + '\n' + r'\hspace{0.05\textwidth}'.join(row2))
    if 'firma' in campos:
        lines.append(r'\\[8mm]\textbf{Firma:} \hspace{8cm}\vspace{4mm}')
    if not lines:
        return ''
    body = '\n'.join(lines)
    return (
        r'\begin{tcolorbox}[colback=white,colframe=secundario,arc=2mm,boxrule=1pt,'
        r'left=5mm,right=5mm,top=3mm,bottom=3mm]' '\n'
        + body + '\n'
        r'\end{tcolorbox}' '\n'
        r'\vspace{3mm}' '\n'
    )


def _gen_seccion_info(pts: str, penalizacion: str) -> str:
    r"""Combina puntos y penalización en una cadena para \seccionexamen."""
    parts = []
    if pts and str(pts).strip():
        parts.append(f"{pts} pts")
    if penalizacion and str(penalizacion).strip() not in ('', 'Sin penalización'):
        parts.append(f"pen: {penalizacion}")
    return ' · '.join(parts)


# --- UTILIDADES LATEX ---
def _escape_latex(text):
    """Escapa caracteres especiales de LaTeX, preservando secciones math ($...$, $$...$$)."""
    text = str(text)
    # Extraer secciones math para protegerlas del escape
    import re as _re
    _math_parts = []
    def _save_math(m):
        _math_parts.append(m.group(0))
        return f'\x00MATH{len(_math_parts)-1}\x00'
    # Proteger $$...$$ primero, luego $...$
    text = _re.sub(r'\$\$.+?\$\$', _save_math, text, flags=_re.DOTALL)
    text = _re.sub(r'\$.+?\$', _save_math, text)
    # Escapar texto plano
    text = text.replace('\\', r'\textbackslash{}')
    for char in ['&', '%', '#', '_', '{', '}']:
        text = text.replace(char, '\\' + char)
    text = text.replace('~', r'\textasciitilde{}')
    text = text.replace('^', r'\textasciicircum{}')
    # Restaurar secciones math
    for i, mp in enumerate(_math_parts):
        text = text.replace(f'\x00MATH{i}\x00', mp)
    return text


def _parse_markdown_runs(paragraph, text, font_name='Calibri', font_size=None):
    """Añade runs con formato bold/italic al párrafo parseando **bold** y *italic* markdown."""
    import re as _re
    from docx.shared import Pt
    parts = _re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', str(text))
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)
    for run in paragraph.runs:
        if font_name: run.font.name = font_name
        if font_size: run.font.size = Pt(font_size)


def _markdown_to_latex(text):
    """Convierte **bold** y *italic* markdown a LaTeX \\textbf{}/\\textit{}. Escapa el resto."""
    import re as _re
    parts = _re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', str(text))
    result = ""
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            result += r'\textbf{' + _escape_latex(part[2:-2]) + '}'
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            result += r'\textit{' + _escape_latex(part[1:-1]) + '}'
        else:
            result += _escape_latex(part)
    return result


# --- EXPORTAR ---
def generar_master_examen(pool, num_modelos, cfg):
    master = []
    letras_version = ['A','B','C','D','E','F']
    for m in range(1, num_modelos+1):
        ex = [p.copy() for p in pool]
        if cfg.get('barajar_preguntas', True): random.shuffle(ex)
        for idx, p in enumerate(ex):
            ops = p.get('opciones_visibles', p.get('opciones_list'))[:]
            correcta_orig_letra = p.get('letra_correcta', 'A')
            idx_orig = {'A':0, 'B':1, 'C':2, 'D':3}.get(correcta_orig_letra, 0)
            txt_correcta = p.get('opciones_list')[idx_orig]
            anclado = False
            frases_anclaje = ["todas las anteriores", "ninguna de las anteriores", "ambas son", "son correctas", "son falsas"]
            if cfg.get('frases_anclaje_extra'):
                frases_anclaje.extend([x.lower().strip() for x in cfg['frases_anclaje_extra'].split(',')])
            for op in ops:
                if any(f in str(op).lower() for f in frases_anclaje): anclado = True
            if cfg.get('barajar_respuestas', True) and not anclado:
                random.shuffle(ops)
            try:
                new_idx = ops.index(txt_correcta)
                new_letra = ['A','B','C','D'][new_idx]
            except: new_letra = 'A'
            p['opciones_finales'] = ops; p['letra_final'] = new_letra; p['num'] = idx + 1
        letra_v = letras_version[m-1] if m <= len(letras_version) else str(m)
        master.append({'modelo': m, 'letra_version': letra_v, 'preguntas': ex})
    return master

def exportar_archivos_csv(master, ruta, nombre):
    # 1. CLAVES (Formato Vertical Lector)
    rows_claves = []
    for m in master:
        version = m['letra_version']
        for p in m['preguntas']:
            rows_claves.append({
                'Pregunta': p['num'], 'Versión': version, 'Respuesta': p['letra_final'],
                'Puntos': '1,00', 'Vínculo': '0', 'Penalización': '0,33',
                'Ponderación': '0,00', 'Máx.P.Abiertas': '0,00'
            })
    pd.DataFrame(rows_claves).to_csv(os.path.join(ruta, f"{nombre}_CLAVES.csv"), index=False, encoding='utf-8-sig', sep=';')

    # 2. METADATA (Información para el profesor)
    rows_meta = []
    for m in master:
        for p in m['preguntas']:
            rows_meta.append({
                'Modelo': m['modelo'], 'Version': m['letra_version'], 'Num_Examen': p['num'],
                'ID_BaseDatos': p['ID_Pregunta'], 'Bloque': p.get('bloque',''), 'Tema': p.get('Tema',''),
                'Dificultad': p.get('dificultad',''), 'Enunciado_Inicio': str(p['enunciado'])[:50]
            })
    pd.DataFrame(rows_meta).to_csv(os.path.join(ruta, f"{nombre}_METADATA.csv"), index=False, encoding='utf-8-sig', sep=';')

def generar_latex(master, ruta, nombre, cfg, modo_solucion=False):
    sufijo = "_SOL" if modo_solucion else ""
    plantilla_tex = ""
    if cfg.get('plantilla_tex_path'):
        try:
            with open(cfg['plantilla_tex_path'], 'r', encoding='utf-8') as f: plantilla_tex = f.read()
        except (FileNotFoundError, PermissionError, UnicodeDecodeError): pass
    _DEFAULT_TEX = r"""\documentclass[a4paper,12pt]{exam}
\usepackage[utf8]{inputenc}
\usepackage[T1]{fontenc}
\usepackage[spanish]{babel}
\usepackage{amsmath,amssymb}
\usepackage{geometry}
\geometry{margin=2cm}
\usepackage{xcolor}
\usepackage{graphicx}

% Header / Footer
\pagestyle{headandfoot}
\firstpageheader{[[LOGO]][[INSTITUCION]]}{[[TIPO_EXAMEN]] -- Modelo [[VERSION]]}{[[FECHA]]}
\runningheader{[[TITULO]]}{Modelo [[VERSION]]}{Pág. \thepage/\numpages}
\firstpagefooter{}{}{}
\runningfooter{}{}{}

\begin{document}

\begin{center}
{\Large \textbf{[[TITULO]]}} \\[6pt]
{\large [[TIPO_EXAMEN]] -- Modelo [[VERSION]]} \\[4pt]
[[INSTITUCION]] \hfill [[FECHA]] \hfill Tiempo: [[TIEMPO]]
\end{center}

\vspace{2mm}
\noindent\textbf{Nombre:} \hrulefill \hspace{1cm} \textbf{Grupo:} \hrulefill

\vspace{3mm}
\noindent\textit{[[INSTRUCCIONES]]}

[[INFO_FUNDAMENTALES]]

[[FUNDAMENTALES]]

[[INFO_TEST]]

\begin{questions}
[[PREGUNTAS]]
\end{questions}

\end{document}"""

    if not plantilla_tex:
        plantilla_tex = _DEFAULT_TEX

    for m in master:
        tex = plantilla_tex
        tex = tex.replace('[[TIPO_EXAMEN]]', _escape_latex(cfg.get('tipo_examen','')))
        tex = tex.replace('[[INSTITUCION]]', _escape_latex(cfg.get('entidad','')))
        tex = tex.replace('[[TITULO]]', _escape_latex(cfg.get('titulo_asignatura','')))
        tex = tex.replace('[[FECHA]]', _escape_latex(cfg.get('fecha','')))
        tex = tex.replace('[[TIEMPO]]', _escape_latex(cfg.get('tiempo','')))
        tex = tex.replace('[[VERSION]]', m['letra_version'])
        tex = tex.replace('[[INSTRUCCIONES]]', _escape_latex(cfg.get('instr_gen','')))
        tex = tex.replace('[[INFO_FUNDAMENTALES]]', _escape_latex(cfg.get('info_fund','')))
        tex = tex.replace('[[INFO_TEST]]', _escape_latex(cfg.get('info_test','')))
        logo_path = cfg.get('logo_path', '')
        if logo_path and os.path.isfile(logo_path):
            logo_tex = r"\includegraphics[height=1.2cm]{" + logo_path.replace('\\', '/') + r"} "
        else:
            logo_tex = ""
        tex = tex.replace('[[LOGO]]', logo_tex)

        bloque_fund = ""
        if cfg.get('fundamentales_data'):
            bloque_fund = r"\begin{enumerate}" + "\n"
            for c in cfg['fundamentales_data']:
                bloque_fund += r"\item " + _markdown_to_latex(c['txt']) + r" (" + str(c['pts']) + " pts)\n"
                esp = c.get('espacio','Automático'); h = "5cm"
                if "10" in esp: h="8cm"
                elif "Media" in esp: h="12cm"
                elif "Cara" in esp: h="18cm"
                if modo_solucion: bloque_fund += r"\par \textit{[Espacio (" + esp + r")]} \vspace{1cm}" + "\n"
                else: bloque_fund += r"\par \framebox{\begin{minipage}[t]["+h+r"]{\linewidth} \end{minipage}} \vspace{0.5cm}" + "\n"
            bloque_fund += r"\end{enumerate}" + "\n"
        tex = tex.replace('[[FUNDAMENTALES]]', bloque_fund)

        bloque_test = ""
        for p in m['preguntas']:
            bloque_test += r"\question " + _escape_latex(p['enunciado']) + "\n"
            ops = p['opciones_finales']; letra_corr = p['letra_final']; idx_corr = {'A':0,'B':1,'C':2,'D':3}.get(letra_corr, 0)
            bloque_test += r"\begin{choices}" + "\n"
            for i, l in enumerate(['a','b','c','d']):
                txt_op = _escape_latex(ops[i])
                if modo_solucion and i == idx_corr:
                    if cfg.get('sol_negrita'): txt_op = r"\textbf{" + txt_op + "}"
                    if cfg.get('sol_rojo'): txt_op = r"\textcolor{red}{" + txt_op + "}"
                    if cfg.get('sol_ast'): txt_op = txt_op + " *"
                    bloque_test += r"\CorrectChoice " + txt_op + "\n"
                else:
                    bloque_test += r"\choice " + txt_op + "\n"
            bloque_test += r"\end{choices}" + "\n"
        tex = tex.replace('[[PREGUNTAS]]', bloque_test)
        
        with open(os.path.join(ruta, f"{nombre}_MOD{m['letra_version']}{sufijo}.tex"), "w", encoding='utf-8') as f: f.write(tex)

def _insert_paragraph_after(paragraph, text=""):
    """Inserta un párrafo nuevo inmediatamente después del párrafo dado. Devuelve el nuevo Paragraph."""
    from docx.oxml.ns import qn
    new_p = paragraph._element.makeelement(qn('w:p'), {})
    paragraph._element.addnext(new_p)
    new_para = paragraph.__class__(new_p, paragraph._parent)
    if text:
        new_para.text = text
    return new_para

def _insert_table_after(paragraph, rows, cols, doc):
    """Inserta una tabla después del párrafo dado. Devuelve la tabla."""
    tbl = doc.add_table(rows=rows, cols=cols)
    # Mover la tabla XML justo después del párrafo
    paragraph._element.addnext(tbl._tbl)
    return tbl

def _replace_in_paragraph(paragraph, key, val):
    """Reemplaza placeholder en un párrafo respetando los runs y su formato."""
    full = paragraph.text
    if key not in full:
        return
    # Caso simple: placeholder en un solo run
    for run in paragraph.runs:
        if key in run.text:
            run.text = run.text.replace(key, val)
            return
    # Caso complejo: placeholder dividido entre runs - reconstruir
    combined = ''.join(r.text for r in paragraph.runs)
    idx = combined.find(key)
    if idx < 0:
        return
    # Localizar qué runs abarca el placeholder
    pos = 0
    start_run = end_run = 0
    start_offset = end_offset = 0
    for ri, run in enumerate(paragraph.runs):
        rlen = len(run.text)
        if pos + rlen > idx and start_run == 0 and pos <= idx:
            start_run = ri
            start_offset = idx - pos
        if pos + rlen >= idx + len(key):
            end_run = ri
            end_offset = idx + len(key) - pos
            break
        pos += rlen
    # Insertar valor en el primer run, limpiar los intermedios
    runs = paragraph.runs
    runs[start_run].text = runs[start_run].text[:start_offset] + val + runs[end_run].text[end_offset:]
    for ri in range(start_run + 1, end_run + 1):
        runs[ri].text = ''

def _setup_word_styles(doc, cfg, version):
    """Configura documento Word moderno: banner título con fondo, logo, caja alumno con encabezado."""
    from docx.shared import Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    scheme   = cfg.get('color_scheme', 'azul')
    cs       = _WORD_COLOR_SCHEMES.get(scheme, _WORD_COLOR_SCHEMES['azul'])
    pri      = RGBColor(*cs['primary'])
    sec      = RGBColor(*cs['secondary'])
    pri_hex  = '%02X%02X%02X' % cs['primary']
    sec_hex  = '%02X%02X%02X' % cs['secondary']
    bg_hex   = '%02X%02X%02X' % cs['bg']
    white    = RGBColor(255, 255, 255)

    # Colores de texto claro para usar sobre fondo de color primario
    _lite = {'azul': (RGBColor(189,215,238), RGBColor(214,232,248)),
             'ucm':  (RGBColor(235,200,205), RGBColor(248,225,230)),
             'byn':  (RGBColor(200,200,200), RGBColor(225,225,225))}
    lite1, lite2 = _lite.get(scheme, _lite['azul'])   # lite1=institución, lite2=tipo/modelo

    font_key   = cfg.get('tipografia', 'cm')
    font_name  = _WORD_FONTS.get(font_key, 'Calibri')
    font_size  = int(cfg.get('font_size', 12))
    linespread = float(cfg.get('linespread', 1.0))
    logo_path  = cfg.get('logo_path', '')
    campos_alumno = cfg.get('campos_alumno', ['nombre', 'dni', 'grupo', 'firma'])

    # ── Helpers ────────────────────────────────────────────────────────────────
    def _shd_cell(cell, hex_col):
        tcPr = cell._tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_col)
        tcPr.append(shd)

    def _no_borders_table(table):
        tbl   = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblBd = OxmlElement('w:tblBorders')
        for bn in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            b = OxmlElement(f'w:{bn}')
            b.set(qn('w:val'), 'none')
            tblBd.append(b)
        tblPr.append(tblBd)

    def _cell_font(run, size=None, bold=False, color=None):
        run.font.name = font_name
        run.font.size = Pt(size or font_size)
        run.font.bold = bold
        if color: run.font.color.rgb = color

    # ── Márgenes ──────────────────────────────────────────────────────────────
    for sec_obj in doc.sections:
        sec_obj.top_margin      = Cm(2.5)
        sec_obj.bottom_margin   = Cm(2.5)
        sec_obj.left_margin     = Cm(2.5)
        sec_obj.right_margin    = Cm(2.5)
        sec_obj.header_distance = Cm(1.2)

    # ── Fuente por defecto ─────────────────────────────────────────────────────
    doc.styles['Normal'].font.name = font_name
    doc.styles['Normal'].font.size = Pt(font_size)

    # ── Header Word (texto pequeño, separador inferior) ───────────────────────
    header = doc.sections[0].header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.clear()
    parts_h = []
    if cfg.get('entidad', ''):           parts_h.append(cfg['entidad'])
    if cfg.get('titulo_asignatura', ''): parts_h.append(cfg['titulo_asignatura'])
    parts_h.append(f"Modelo {version}")
    rh = hp.add_run("  ·  ".join(parts_h))
    rh.font.name = font_name; rh.font.size = Pt(8); rh.font.color.rgb = sec
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr_h  = hp._p.get_or_add_pPr()
    pBdr_h = OxmlElement('w:pBdr')
    bot_h  = OxmlElement('w:bottom')
    bot_h.set(qn('w:val'), 'single'); bot_h.set(qn('w:sz'), '4')
    bot_h.set(qn('w:space'), '1');    bot_h.set(qn('w:color'), sec_hex)
    pBdr_h.append(bot_h); pPr_h.append(pBdr_h)

    # ── Footer (nº de página centrado) ────────────────────────────────────────
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    r1 = fp.add_run()
    fld_b = OxmlElement('w:fldChar'); fld_b.set(qn('w:fldCharType'), 'begin')
    r1._r.append(fld_b)
    r2 = fp.add_run()
    instr_el = OxmlElement('w:instrText'); instr_el.text = ' PAGE '
    r2._r.append(instr_el)
    r3 = fp.add_run()
    fld_e = OxmlElement('w:fldChar'); fld_e.set(qn('w:fldCharType'), 'end')
    r3._r.append(fld_e)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in fp.runs: r.font.size = Pt(8); r.font.color.rgb = sec

    # ── BANNER TÍTULO (tabla sin bordes con fondo de color primario) ───────────
    has_logo = bool(logo_path and os.path.isfile(logo_path))
    tbl_ban  = doc.add_table(rows=1, cols=2 if has_logo else 1)
    _no_borders_table(tbl_ban)

    if has_logo:
        logo_cell = tbl_ban.rows[0].cells[0]
        _shd_cell(logo_cell, pri_hex)
        # Fijar ancho de columna logo (~3 cm = 1701 twips)
        tcPr_l = logo_cell._tc.get_or_add_tcPr()
        tcW_l  = OxmlElement('w:tcW')
        tcW_l.set(qn('w:w'), '1701'); tcW_l.set(qn('w:type'), 'dxa')
        tcPr_l.append(tcW_l)
        p_logo = logo_cell.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_before = Pt(10)
        p_logo.paragraph_format.space_after  = Pt(10)
        try:
            p_logo.add_run().add_picture(logo_path, height=Cm(1.5))
        except Exception:
            pass  # Ignorar si el logo no se puede cargar

    tit_cell = tbl_ban.rows[0].cells[-1]
    _shd_cell(tit_cell, pri_hex)

    # Línea 1: institución (letra pequeña, color claro)
    p0 = tit_cell.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(10)
    p0.paragraph_format.space_after  = Pt(2)
    if cfg.get('entidad', ''):
        r0 = p0.add_run(cfg['entidad'].upper())
        _cell_font(r0, size=font_size - 2, color=lite1)

    # Línea 2: asignatura (grande, negrita, blanca)
    p1 = tit_cell.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(2)
    p1.paragraph_format.space_after  = Pt(4)
    r1b = p1.add_run(cfg.get('titulo_asignatura', ''))
    _cell_font(r1b, size=font_size + 5, bold=True, color=white)

    # Línea 3: tipo + modelo (mediana, color suave)
    sub_parts = []
    if cfg.get('tipo_examen', ''): sub_parts.append(cfg['tipo_examen'])
    sub_parts.append(f"Modelo {version}")
    p2 = tit_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(10)
    r2b = p2.add_run("  —  ".join(sub_parts))
    _cell_font(r2b, size=font_size, color=lite2)

    # ── FRANJA INFO (fecha · tiempo sobre fondo suave) ─────────────────────────
    info_parts = []
    if cfg.get('fecha', ''):  info_parts.append(cfg['fecha'])
    if cfg.get('tiempo', ''): info_parts.append(f"Tiempo: {cfg['tiempo']}")
    if info_parts:
        p_info = doc.add_paragraph()
        p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_info.paragraph_format.space_before = Pt(0)
        p_info.paragraph_format.space_after  = Pt(4)
        pPr_i = p_info._p.get_or_add_pPr()
        shd_i = OxmlElement('w:shd')
        shd_i.set(qn('w:val'), 'clear'); shd_i.set(qn('w:color'), 'auto')
        shd_i.set(qn('w:fill'), bg_hex)
        pPr_i.append(shd_i)
        r_info = p_info.add_run("     ·     ".join(info_parts))
        r_info.font.name = font_name; r_info.font.size = Pt(font_size - 1)
        r_info.font.color.rgb = sec

    # ── DIVISORIA GRUESA ──────────────────────────────────────────────────────
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_before = Pt(2)
    p_div.paragraph_format.space_after  = Pt(6)
    pPr_d  = p_div._p.get_or_add_pPr()
    pBdr_d = OxmlElement('w:pBdr')
    bot_d  = OxmlElement('w:bottom')
    bot_d.set(qn('w:val'), 'single'); bot_d.set(qn('w:sz'), '12')
    bot_d.set(qn('w:space'), '1');    bot_d.set(qn('w:color'), pri_hex)
    pBdr_d.append(bot_d); pPr_d.append(pBdr_d)

    # ── INSTRUCCIONES (fondo suave, borde izquierdo coloreado) ────────────────
    instr_txt = cfg.get('instr_gen', '').strip()
    if instr_txt:
        p_ins = doc.add_paragraph()
        pPr_ins = p_ins._p.get_or_add_pPr()
        shd_ins = OxmlElement('w:shd')
        shd_ins.set(qn('w:val'), 'clear'); shd_ins.set(qn('w:color'), 'auto')
        shd_ins.set(qn('w:fill'), bg_hex)
        pPr_ins.append(shd_ins)
        pBdr_ins = OxmlElement('w:pBdr')
        left_b   = OxmlElement('w:left')
        left_b.set(qn('w:val'), 'single'); left_b.set(qn('w:sz'), '18')
        left_b.set(qn('w:space'), '4');    left_b.set(qn('w:color'), pri_hex)
        pBdr_ins.append(left_b); pPr_ins.append(pBdr_ins)
        r_ins = p_ins.add_run(instr_txt)
        r_ins.italic = True; r_ins.font.size = Pt(font_size - 1); r_ins.font.name = font_name
        p_ins.paragraph_format.space_before = Pt(0)
        p_ins.paragraph_format.space_after  = Pt(8)

    # ── CAJA DATOS DEL ALUMNO (después de instrucciones, con encabezado) ──────
    if campos_alumno:
        tbl_a = doc.add_table(rows=1, cols=2)
        tbl_a.style = 'Table Grid'

        # Fila 0: encabezado "DATOS DEL ALUMNO" (fondo primario, texto blanco)
        hdr_cell = tbl_a.rows[0].cells[0].merge(tbl_a.rows[0].cells[1])
        _shd_cell(hdr_cell, pri_hex)
        p_hdr = hdr_cell.paragraphs[0]
        p_hdr.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_hdr.paragraph_format.space_before = Pt(3)
        p_hdr.paragraph_format.space_after  = Pt(3)
        r_hdr = p_hdr.add_run("DATOS DEL ALUMNO")
        _cell_font(r_hdr, size=font_size - 1, bold=True, color=white)

        # Fila nombre (merged, ancho completo)
        if 'nombre' in campos_alumno:
            row_n  = tbl_a.add_row()
            cell_n = row_n.cells[0].merge(row_n.cells[1])
            p_n    = cell_n.paragraphs[0]
            r_nl   = p_n.add_run("Nombre y Apellidos: ")
            _cell_font(r_nl, bold=True)
            p_n.add_run("_" * 58)
            p_n.paragraph_format.space_before = Pt(4)
            p_n.paragraph_format.space_after  = Pt(4)

        # Filas para dni / grupo / firma (2 por fila)
        extra  = [f for f in ('dni', 'grupo', 'firma') if f in campos_alumno]
        labels = {'dni': 'DNI/NIU', 'grupo': 'Grupo', 'firma': 'Firma'}
        for i in range(0, len(extra), 2):
            chunk = extra[i:i+2]
            row_e = tbl_a.add_row()
            if len(chunk) == 1:
                cell_e = row_e.cells[0].merge(row_e.cells[1])
                p_e    = cell_e.paragraphs[0]
                r_e    = p_e.add_run(f"{labels[chunk[0]]}: ")
                _cell_font(r_e, bold=True)
                p_e.add_run("_" * 45)
            else:
                for ci, fname in enumerate(chunk):
                    p_e = row_e.cells[ci].paragraphs[0]
                    r_e = p_e.add_run(f"{labels[fname]}: ")
                    _cell_font(r_e, bold=True)
                    p_e.add_run("_" * 22)
                    p_e.paragraph_format.space_before = Pt(4)
                    p_e.paragraph_format.space_after  = Pt(4)
            if len(chunk) == 1:
                p_e.paragraph_format.space_before = Pt(4)
                p_e.paragraph_format.space_after  = Pt(4)

        doc.add_paragraph()

    # ── Aplicar interlineado a todos los párrafos ya añadidos ─────────────────
    if linespread != 1.0:
        line_val = str(int(linespread * 240))
        for p in doc.paragraphs:
            pPr_s = p._p.get_or_add_pPr()
            spng  = OxmlElement('w:spacing')
            spng.set(qn('w:line'), line_val)
            spng.set(qn('w:lineRule'), 'auto')
            pPr_s.append(spng)

    doc.add_paragraph("[[FUNDAMENTALES]]")
    doc.add_paragraph("[[PREGUNTAS]]")

def rellenar_plantilla_word(master, ruta, nombre, cfg, tpl_path=None, modo_solucion=False):
    for m in master:
        doc = Document(tpl_path) if tpl_path else Document()
        sufijo = "_SOL" if modo_solucion else ""
        if not tpl_path:
            _setup_word_styles(doc, cfg, m['letra_version'])

        replacements = {'[[TIPO_EXAMEN]]': cfg.get('tipo_examen',''), '[[INSTITUCION]]': cfg.get('entidad',''),
            '[[TITULO]]': cfg.get('titulo_asignatura',''), '[[FECHA]]': cfg.get('fecha',''),
            '[[TIEMPO]]': cfg.get('tiempo',''), '[[VERSION]]': m['letra_version'],
            '[[INSTRUCCIONES]]': cfg.get('instr_gen',''), '[[INFO_FUNDAMENTALES]]': cfg.get('info_fund',''),
            '[[INFO_TEST]]': cfg.get('info_test','')}

        for p in doc.paragraphs:
            for key, val in replacements.items():
                _replace_in_paragraph(p, key, val)

        for p in list(doc.paragraphs):
            if '[[FUNDAMENTALES]]' in p.text:
                p.text = ""
                insert_after = p
                if cfg.get('fundamentales_data'):
                    for c in cfg['fundamentales_data']:
                        p_fund = _insert_paragraph_after(insert_after)
                        _parse_markdown_runs(p_fund, c['txt'], font_name='Calibri', font_size=11)
                        r_pts = p_fund.add_run(f"  ({c['pts']} pts)")
                        r_pts.italic = True; r_pts.font.name = 'Calibri'; r_pts.font.size = Pt(10)
                        insert_after = p_fund
                        if not modo_solucion:
                            esp = c.get('espacio','Automático')
                            tbl = _insert_table_after(insert_after, rows=1, cols=1, doc=doc)
                            tbl.style = 'Table Grid'
                            lines = 3 if "5" in esp else 6 if "10" in esp else 10
                            for _ in range(lines): tbl.rows[0].cells[0].add_paragraph()
                            # Avanzar insert_after al párrafo siguiente a la tabla
                            # Creamos un párrafo separador tras la tabla
                            sep_p_el = tbl._tbl.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p', {})
                            tbl._tbl.addnext(sep_p_el)
                            from docx.text.paragraph import Paragraph
                            insert_after = Paragraph(sep_p_el, doc.element.body)

            if '[[PREGUNTAS]]' in p.text:
                p.text = ""
                insert_after = p
                for preg in m['preguntas']:
                    p_enun = _insert_paragraph_after(insert_after)
                    r_num = p_enun.add_run(f"{preg['num']}. {preg['enunciado']}")
                    r_num.bold = True; r_num.font.name = 'Calibri'; r_num.font.size = Pt(11)
                    insert_after = p_enun
                    ops = preg['opciones_finales']; letra_corr = preg['letra_final']; idx_corr = {'A':0,'B':1,'C':2,'D':3}.get(letra_corr, 0)
                    for i, l in enumerate(['a','b','c','d']):
                        p_op = _insert_paragraph_after(insert_after, f"    {l}) {ops[i]}")
                        if modo_solucion and i == idx_corr:
                            if cfg.get('sol_negrita'): p_op.runs[0].bold = True
                            if cfg.get('sol_rojo'): p_op.runs[0].font.color.rgb = RGBColor(255, 0, 0)
                            if cfg.get('sol_ast'): p_op.add_run(" *")
                        insert_after = p_op

        doc.save(os.path.join(ruta, f"{nombre}_MOD{m['letra_version']}{sufijo}.docx"))


# ── Variantes en memoria (para Streamlit Cloud / descarga directa) ─────────────

def exportar_csv_bytes(master, nombre) -> dict:
    """Genera los CSVs en memoria. Retorna {'claves': bytes, 'metadata': bytes}."""
    import io as _io
    rows_claves, rows_meta = [], []
    for m in master:
        version = m['letra_version']
        for p in m['preguntas']:
            rows_claves.append({
                'Pregunta': p['num'], 'Versión': version, 'Respuesta': p['letra_final'],
                'Puntos': '1,00', 'Vínculo': '0', 'Penalización': '0,33',
                'Ponderación': '0,00', 'Máx.P.Abiertas': '0,00'
            })
            rows_meta.append({
                'Modelo': m['modelo'], 'Version': version, 'Num_Examen': p['num'],
                'ID_BaseDatos': p['ID_Pregunta'], 'Bloque': p.get('bloque', ''),
                'Tema': p.get('Tema', ''), 'Dificultad': p.get('dificultad', ''),
                'Enunciado_Inicio': str(p['enunciado'])[:50]
            })
    buf_c, buf_m = _io.BytesIO(), _io.BytesIO()
    pd.DataFrame(rows_claves).to_csv(buf_c, index=False, encoding='utf-8-sig', sep=';')
    pd.DataFrame(rows_meta).to_csv(buf_m, index=False, encoding='utf-8-sig', sep=';')
    return {'claves': buf_c.getvalue(), 'metadata': buf_m.getvalue()}


def generar_latex_strings(master, nombre, cfg, modo_solucion=False) -> dict:
    """Genera los .tex en memoria. Retorna {letra_version: str_tex}."""
    # Si hay plantilla personalizada, usarla directamente
    plantilla_tex = ""
    tpl_bytes = cfg.get('plantilla_tex_bytes')
    if tpl_bytes:
        try:
            plantilla_tex = tpl_bytes.decode('utf-8')
        except Exception:
            pass
    elif cfg.get('plantilla_tex_path'):
        try:
            with open(cfg['plantilla_tex_path'], 'r', encoding='utf-8') as f:
                plantilla_tex = f.read()
        except Exception:
            pass

    # ── Parámetros de estilo ─────────────────────────────────────────────────
    font_size = int(cfg.get('font_size', 12))
    doc_class = 'extarticle' if font_size > 12 else 'article'
    linespread = cfg.get('linespread', 1.0)
    linespread_cmd = (
        f'\\usepackage{{setspace}}\\setstretch{{{linespread}}}'
        if linespread != 1.0 else ''
    )
    campos_alumno = cfg.get('campos_alumno', ['nombre', 'dni', 'grupo', 'firma'])
    opciones_cols = cfg.get('opciones_cols', 1)

    # Puntos/penalización → info para \seccionexamen
    pts_fund = str(cfg.get('pts_fund', '')).strip()
    pts_test = str(cfg.get('pts_test', '')).strip()
    penalizacion = str(cfg.get('penalizacion', '')).strip()
    info_fund_sec = _gen_seccion_info(pts_fund, '')
    info_test_sec = _gen_seccion_info(pts_test, penalizacion)

    # Identificador de versión adaptada (se añade al tipo_examen si presente)
    adapt_id = str(cfg.get('adaptada_id', '')).strip()
    tipo_raw = cfg.get('tipo_examen', '')
    if adapt_id:
        tipo_raw = f"{tipo_raw} — {adapt_id}" if tipo_raw else adapt_id

    _DEFAULT_TEX = r"""\documentclass[a4paper,[[FONTSIZE]]pt]{[[DOCTYPE]]}
\usepackage{estilo_examen_moderno_v2}
[[LINESPREAD_CMD]]
\tipoexamen{[[TIPO_EXAMEN]]}
\institucion{[[INSTITUCION]]}
\asignatura{[[TITULO]]}
\fecha{[[FECHA]]}
\duracion{[[TIEMPO]]}
\version{[[VERSION]]}
[[LOGO_CMD]]
\begin{document}
\encabezadoprofesional
[[CAJA_ALUMNO]]
[[INSTRUCCIONES_TEX]]
[[BLOQUE_FUND]]
[[BLOQUE_TEST]]
\end{document}"""

    if not plantilla_tex:
        plantilla_tex = _DEFAULT_TEX

    result = {}
    for m in master:
        tex = plantilla_tex
        tex = tex.replace('[[FONTSIZE]]',    str(font_size))
        tex = tex.replace('[[DOCTYPE]]',     doc_class)
        tex = tex.replace('[[LINESPREAD_CMD]]', linespread_cmd)
        tex = tex.replace('[[TIPO_EXAMEN]]', _escape_latex(tipo_raw))
        tex = tex.replace('[[INSTITUCION]]', _escape_latex(cfg.get('entidad', '')))
        tex = tex.replace('[[TITULO]]',      _escape_latex(cfg.get('titulo_asignatura', '')))
        tex = tex.replace('[[FECHA]]',       _escape_latex(cfg.get('fecha', '')))
        tex = tex.replace('[[TIEMPO]]',      _escape_latex(cfg.get('tiempo', '')))
        tex = tex.replace('[[VERSION]]',     m['letra_version'])
        logo_path = cfg.get('logo_path', '')
        tex = tex.replace('[[LOGO_CMD]]', f'\\logo{{{logo_path}}}' if logo_path else '')

        # Caja alumno
        tex = tex.replace('[[CAJA_ALUMNO]]', _gen_caja_alumno_latex(campos_alumno))

        # Instrucciones
        instr = cfg.get('instr_gen', '').strip()
        tex = tex.replace('[[INSTRUCCIONES_TEX]]',
                          '\\instrucciones{' + _escape_latex(instr) + '}\n' if instr else '')

        # ── PARTE I: Desarrollo ──────────────────────────────────────────────
        bloque_fund = ""
        fund_data = cfg.get('fundamentales_data', [])
        if fund_data:
            tit_fund = _escape_latex(cfg.get('titulo_fund', 'PREGUNTAS DE DESARROLLO'))
            bloque_fund += f'\\seccionexamen{{PARTE I --- {tit_fund}}}{{{info_fund_sec}}}\n'
            if cfg.get('info_fund', '').strip():
                bloque_fund += '\\textit{' + _escape_latex(cfg['info_fund']) + '}\n\n'
            bloque_fund += '\\begin{enumerate}\n'
            _esp_map = {'5 líneas': '4cm', '10 líneas': '7cm',
                        'media cara': '11cm', 'cara completa': '20cm'}
            for c in fund_data:
                esp    = c.get('espacio', 'Automático')
                h_base = next((v for k, v in _esp_map.items() if k.lower() in esp.lower()), '6cm')
                # Modo adaptado: aumentar altura si adapt_espacio_extra definido
                extra_pct = cfg.get('adapt_espacio_pct', 0)
                if extra_pct:
                    import re as _re2
                    m2 = _re2.match(r'([\d.]+)cm', h_base)
                    if m2:
                        h_base = f"{float(m2.group(1)) * (1 + extra_pct/100):.1f}cm"
                pts_q = str(c.get('pts', '')).strip()
                pts_str = f' ({pts_q} pts)' if pts_q else ''
                bloque_fund += '\\item \\begin{minipage}[t]{\\linewidth}\n'
                bloque_fund += _markdown_to_latex(c['txt']) + pts_str + '\n'
                if modo_solucion:
                    bloque_fund += f'\\par\\textit{{[Espacio de respuesta ({esp})]}}\n'
                else:
                    bloque_fund += f'\\espaciorespuesta[{h_base}]\n'
                bloque_fund += '\\end{minipage}\n\n'
            bloque_fund += '\\end{enumerate}\n'
        tex = tex.replace('[[BLOQUE_FUND]]', bloque_fund)

        # ── PARTE II: Test ───────────────────────────────────────────────────
        bloque_test = ""
        if m.get('preguntas'):
            num_parte = 'II' if fund_data else 'I'
            tit_test  = _escape_latex(cfg.get('titulo_test', 'PREGUNTAS TEST'))
            bloque_test += f'\\seccionexamen{{PARTE {num_parte} --- {tit_test}}}{{{info_test_sec}}}\n'
            if cfg.get('info_test', '').strip():
                bloque_test += '\\textit{' + _escape_latex(cfg['info_test']) + '}\n\n'
            bloque_test += '\\begin{enumerate}[resume]\n' if fund_data else '\\begin{enumerate}\n'
            for p in m['preguntas']:
                ops       = p['opciones_finales']
                letra_c   = p['letra_final']
                idx_c     = {'A': 0, 'B': 1, 'C': 2, 'D': 3}.get(letra_c, 0)
                bloque_test += '\\item \\begin{minipage}[t]{\\linewidth}\n'
                bloque_test += _markdown_to_latex(p['enunciado']) + '\n'
                # Opciones: 1 columna o 2 columnas
                if opciones_cols == 2:
                    bloque_test += '\\begin{enumerate}[label=\\textcolor{secundario}{\\textbf{\\alph*)}},leftmargin=2em,itemsep=0.3em,topsep=0.2em]\n'
                    for i in range(0, 4, 2):
                        txt0 = _escape_latex(ops[i])
                        txt1 = _escape_latex(ops[i+1]) if i+1 < len(ops) else ''
                        if modo_solucion:
                            if i   == idx_c: txt0 = _fmt_sol_latex(txt0, cfg)
                            if i+1 == idx_c: txt1 = _fmt_sol_latex(txt1, cfg)
                        bloque_test += (
                            f'\\item \\begin{{minipage}}[t]{{0.44\\linewidth}}{txt0}\\end{{minipage}}'
                            f'\\hfill'
                            f'\\begin{{minipage}}[t]{{0.44\\linewidth}}{txt1}\\end{{minipage}}\n'
                        )
                else:
                    bloque_test += '\\begin{enumerate}[label=\\textcolor{secundario}{\\textbf{\\alph*)}},leftmargin=2em,itemsep=0.3em,topsep=0.2em]\n'
                    for i in range(4):
                        txt_op = _escape_latex(ops[i])
                        if modo_solucion and i == idx_c:
                            txt_op = _fmt_sol_latex(txt_op, cfg)
                        bloque_test += f'\\item {txt_op}\n'
                bloque_test += '\\end{enumerate}\n'
                bloque_test += '\\end{minipage}\n\n'
            bloque_test += '\\end{enumerate}\n'
        tex = tex.replace('[[BLOQUE_TEST]]', bloque_test)

        result[m['letra_version']] = tex
    return result


def _fmt_sol_latex(txt: str, cfg: dict) -> str:
    """Aplica marcado de solución a una opción correcta en LaTeX."""
    if cfg.get('sol_negrita'): txt = r'\textbf{' + txt + '}'
    if cfg.get('sol_rojo'):    txt = r'\textcolor{red}{' + txt + '}'
    if cfg.get('sol_ast'):     txt = txt + r' \textbf{*}'
    return txt


def rellenar_plantilla_word_bytes(master, nombre, cfg, tpl_bytes=None, modo_solucion=False) -> dict:
    """Genera los .docx en memoria. Retorna {letra_version: bytes_docx}."""
    import io as _io
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE
    from docx.text.paragraph import Paragraph as _Para

    scheme    = cfg.get('color_scheme', 'azul')
    cs        = _WORD_COLOR_SCHEMES.get(scheme, _WORD_COLOR_SCHEMES['azul'])
    pri_rgb   = RGBColor(*cs['primary'])
    sec_rgb   = RGBColor(*cs['secondary'])
    bg_hex    = '%02X%02X%02X' % cs['bg']
    pri_hex   = '%02X%02X%02X' % cs['primary']

    font_key   = cfg.get('tipografia', 'cm')
    font_name  = _WORD_FONTS.get(font_key, 'Calibri')
    font_size  = int(cfg.get('font_size', 12))
    linespread = float(cfg.get('linespread', 1.0))
    opciones_cols = cfg.get('opciones_cols', 1)

    pts_fund    = str(cfg.get('pts_fund', '')).strip()
    pts_test    = str(cfg.get('pts_test', '')).strip()
    penalizacion = str(cfg.get('penalizacion', '')).strip()
    fund_data   = cfg.get('fundamentales_data', [])

    adapt_id    = str(cfg.get('adaptada_id', '')).strip()
    tipo_raw    = cfg.get('tipo_examen', '')
    if adapt_id:
        tipo_raw = f"{tipo_raw} — {adapt_id}" if tipo_raw else adapt_id

    def _set_linespread(para):
        if linespread == 1.0:
            return
        pPr_s = para._p.get_or_add_pPr()
        sp    = _OE('w:spacing')
        sp.set(_qn('w:line'), str(int(linespread * 240)))
        sp.set(_qn('w:lineRule'), 'auto')
        pPr_s.append(sp)

    def _keep_together(para):
        pPr = para._p.get_or_add_pPr()
        kwn = _OE('w:keepNext')
        pPr.append(kwn)
        _set_linespread(para)

    def _add_section_word(doc_obj, titulo, info_str, insert_after_p=None):
        """Añade separador de sección con fondo de color primario."""
        p_sec = _insert_paragraph_after(insert_after_p) if insert_after_p else doc_obj.add_paragraph()
        pPr_sec = p_sec._p.get_or_add_pPr()
        shd_sec = _OE('w:shd')
        shd_sec.set(_qn('w:val'), 'clear'); shd_sec.set(_qn('w:color'), 'auto')
        shd_sec.set(_qn('w:fill'), pri_hex)
        pPr_sec.append(shd_sec)
        p_sec.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r_sec = p_sec.add_run(titulo)
        r_sec.bold = True; r_sec.font.name = font_name
        r_sec.font.size = Pt(font_size); r_sec.font.color.rgb = RGBColor(255, 255, 255)
        if info_str:
            r_info = p_sec.add_run(f"  [{info_str}]")
            r_info.bold = False; r_info.font.name = font_name
            r_info.font.size = Pt(font_size - 1); r_info.font.color.rgb = RGBColor(220, 220, 220)
        p_sec.paragraph_format.space_before = Pt(10)
        p_sec.paragraph_format.space_after  = Pt(4)
        return p_sec

    result = {}
    for m in master:
        doc = Document(_io.BytesIO(tpl_bytes)) if tpl_bytes else Document()
        if not tpl_bytes:
            cfg_for_setup       = dict(cfg)
            cfg_for_setup['tipo_examen'] = tipo_raw
            _setup_word_styles(doc, cfg_for_setup, m['letra_version'])
        else:
            # Plantilla personalizada: hacer reemplazos de texto
            replacements = {
                '[[TIPO_EXAMEN]]': tipo_raw,
                '[[INSTITUCION]]': cfg.get('entidad', ''),
                '[[TITULO]]':      cfg.get('titulo_asignatura', ''),
                '[[FECHA]]':       cfg.get('fecha', ''),
                '[[TIEMPO]]':      cfg.get('tiempo', ''),
                '[[VERSION]]':     m['letra_version'],
                '[[INSTRUCCIONES]]': cfg.get('instr_gen', ''),
            }
            for p in doc.paragraphs:
                for key, val in replacements.items():
                    _replace_in_paragraph(p, key, val)

        # ── PARTE I: Desarrollo ───────────────────────────────────────────────
        for p in list(doc.paragraphs):
            if '[[FUNDAMENTALES]]' not in p.text:
                continue
            p.text = ""
            ins_after = p
            if fund_data:
                info_fund_str = _gen_seccion_info(pts_fund, '')
                tit_f = cfg.get('titulo_fund', 'PREGUNTAS DE DESARROLLO')
                num_f = 'PARTE I'
                ins_after = _add_section_word(doc, f"{num_f} — {tit_f}", info_fund_str, ins_after)
                if cfg.get('info_fund', '').strip():
                    p_hf = _insert_paragraph_after(ins_after)
                    r_hf = p_hf.add_run(cfg['info_fund'])
                    r_hf.italic = True; r_hf.font.name = font_name; r_hf.font.size = Pt(font_size - 1)
                    ins_after = p_hf

                _esp_map = {'5 líneas': 3, '10 líneas': 7, 'media cara': 14, 'cara completa': 28}
                extra_pct = cfg.get('adapt_espacio_pct', 0)

                for ci, c in enumerate(fund_data):
                    esp  = c.get('espacio', 'Automático')
                    pts_q = str(c.get('pts', '')).strip()
                    lines = next((v for k, v in _esp_map.items() if k.lower() in esp.lower()), 8)
                    if extra_pct:
                        lines = int(lines * (1 + extra_pct / 100))

                    p_fund = _insert_paragraph_after(ins_after)
                    r_num_f = p_fund.add_run(f"{ci+1}. ")
                    r_num_f.bold = True; r_num_f.font.name = font_name; r_num_f.font.size = Pt(font_size)
                    _parse_markdown_runs(p_fund, c['txt'], font_name=font_name, font_size=font_size)
                    if pts_q:
                        r_pts = p_fund.add_run(f"  ({pts_q} pts)")
                        r_pts.italic = True; r_pts.font.name = font_name; r_pts.font.size = Pt(font_size - 1)
                    _keep_together(p_fund)
                    ins_after = p_fund

                    if not modo_solucion:
                        tbl = _insert_table_after(ins_after, rows=1, cols=1, doc=doc)
                        tbl.style = 'Table Grid'
                        for _ in range(lines):
                            tbl.rows[0].cells[0].add_paragraph()
                        sep_el = tbl._tbl.makeelement(
                            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p', {})
                        tbl._tbl.addnext(sep_el)
                        ins_after = _Para(sep_el, doc.element.body)
                    else:
                        p_esp = _insert_paragraph_after(ins_after)
                        r_esp = p_esp.add_run(f"[Espacio de respuesta — {esp}]")
                        r_esp.italic = True; r_esp.font.size = Pt(font_size - 2); r_esp.font.name = font_name
                        r_esp.font.color.rgb = RGBColor(*cs['secondary'])
                        ins_after = p_esp

        # ── PARTE II: Test ─────────────────────────────────────────────────────
        for p in list(doc.paragraphs):
            if '[[PREGUNTAS]]' not in p.text:
                continue
            p.text = ""
            ins_after = p
            if not m.get('preguntas'):
                continue
            info_test_str = _gen_seccion_info(pts_test, penalizacion)
            tit_t   = cfg.get('titulo_test', 'PREGUNTAS TEST')
            num_t   = 'II' if fund_data else 'I'
            ins_after = _add_section_word(doc, f"PARTE {num_t} — {tit_t}", info_test_str, ins_after)
            if cfg.get('info_test', '').strip():
                p_ht = _insert_paragraph_after(ins_after)
                r_ht = p_ht.add_run(cfg['info_test'])
                r_ht.italic = True; r_ht.font.name = font_name; r_ht.font.size = Pt(font_size - 1)
                ins_after = p_ht

            base_num = len(fund_data) + 1 if fund_data else 1
            for qi, preg in enumerate(m['preguntas']):
                ops      = preg['opciones_finales']
                letra_c  = preg['letra_final']
                idx_c    = {'A': 0, 'B': 1, 'C': 2, 'D': 3}.get(letra_c, 0)

                p_enun = _insert_paragraph_after(ins_after)
                r_num  = p_enun.add_run(f"{base_num + qi}. ")
                r_num.bold = True; r_num.font.name = font_name; r_num.font.size = Pt(font_size)
                _parse_markdown_runs(p_enun, preg['enunciado'], font_name=font_name, font_size=font_size)
                _keep_together(p_enun)
                ins_after = p_enun

                labels = ['a', 'b', 'c', 'd']
                if opciones_cols == 2:
                    # Dos opciones por fila con tabla
                    tbl_op = _insert_table_after(ins_after, rows=2, cols=2, doc=doc)
                    tbl_op.style = 'Table Grid'
                    for row_i in range(2):
                        for col_i in range(2):
                            oi = row_i * 2 + col_i
                            cell   = tbl_op.rows[row_i].cells[col_i]
                            p_cell = cell.paragraphs[0]
                            p_cell.clear()
                            r_lbl = p_cell.add_run(f"{labels[oi]}) ")
                            r_lbl.font.name = font_name; r_lbl.font.size = Pt(font_size); r_lbl.font.color.rgb = sec_rgb
                            r_txt = p_cell.add_run(ops[oi])
                            r_txt.font.name = font_name; r_txt.font.size = Pt(font_size)
                            if modo_solucion and oi == idx_c:
                                if cfg.get('sol_negrita'): r_txt.bold = True
                                if cfg.get('sol_rojo'):    r_txt.font.color.rgb = RGBColor(255, 0, 0)
                                if cfg.get('sol_ast'):     p_cell.add_run(" *")
                    sep_el2 = tbl_op._tbl.makeelement(
                        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p', {})
                    tbl_op._tbl.addnext(sep_el2)
                    ins_after = _Para(sep_el2, doc.element.body)
                else:
                    for i, lbl in enumerate(labels):
                        p_op = _insert_paragraph_after(ins_after, f"   {lbl}) {ops[i]}")
                        p_op.runs[0].font.name = font_name
                        p_op.runs[0].font.size = Pt(font_size)
                        # Color label
                        p_op.runs[0].font.color.rgb = sec_rgb
                        if modo_solucion and i == idx_c:
                            if cfg.get('sol_negrita'): p_op.runs[0].bold = True
                            if cfg.get('sol_rojo'):    p_op.runs[0].font.color.rgb = RGBColor(255, 0, 0)
                            if cfg.get('sol_ast'):     p_op.add_run(" *")
                        if i < 3:
                            _keep_together(p_op)
                        _set_linespread(p_op)
                        ins_after = p_op

        buf = _io.BytesIO()
        doc.save(buf)
        result[m['letra_version']] = buf.getvalue()
    return result


def generar_zip_bytes(files_dict: dict) -> bytes:
    """Empaqueta archivos en un ZIP en memoria. files_dict: {nombre_archivo: bytes_o_str}."""
    import io as _io, zipfile
    buf = _io.BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for fname, data in files_dict.items():
            if isinstance(data, str):
                data = data.encode('utf-8')
            zf.writestr(fname, data)
    return buf.getvalue()


def cargar_examen_csv(path):
    # Detectar separador leyendo la primera línea
    with open(path, 'r', encoding='utf-8-sig', errors='replace') as f:
        first_line = f.readline()
    sep = ';' if first_line.count(';') >= first_line.count(',') else ','
    df = pd.read_csv(path, sep=sep, encoding='utf-8-sig')
    if 'ID_BaseDatos' in df.columns:
        return df['ID_BaseDatos'].unique().tolist()
    elif 'ID' in df.columns:
        return df['ID'].unique().tolist()
    return []

def exportar_preguntas_json(ids, df, filepath):
    """Exporta una seleccion de preguntas a JSON portable."""
    import json
    pregs = df[df['ID_Pregunta'].isin(ids)].to_dict('records')
    with open(filepath, 'w', encoding='utf-8') as f:
        json.dump({'version': 1, 'preguntas': pregs}, f, ensure_ascii=False, indent=2, default=str)
    return len(pregs)

def importar_preguntas_json(filepath, bloque_destino, df_existing):
    """Importa preguntas desde JSON. Retorna (nuevas_list, duplicados_count)."""
    import json
    with open(filepath, 'r', encoding='utf-8') as f:
        data = json.load(f)
    pregs = data.get('preguntas', [])
    nuevas = []; dupes = 0
    for p in pregs:
        enun = p.get('enunciado','')
        is_dup, _ = check_for_similar_enunciado(enun, df_existing)
        if is_dup: dupes += 1; continue
        tema = p.get('Tema', '1')
        nid, _ = generar_siguiente_id(df_existing, bloque_destino, tema)
        p_new = dict(p)
        p_new['ID_Pregunta'] = nid; p_new['bloque'] = bloque_destino; p_new['usada'] = ''
        nuevas.append(p_new)
        # Update df_existing for next ID generation
        df_existing.loc[len(df_existing)] = p_new
    return nuevas, dupes
